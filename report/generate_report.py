"""
WHO Health IQ - 6-month automated DOCX report generator.

Quantitative, table-heavy report focused on statistical analysis. Claude is
used only for the parts that require interpretation (theme grouping,
executive observations, priority ranking, recommendations) and is asked to
return STRUCTURED JSON so all counts and percentages can be computed in
Python from real data.

Usage
-----
    # Auto-detect current half-year (H1 = Jan-Jun, H2 = Jul-Dec)
    python generate_report.py

    # Explicit half-year
    python generate_report.py --half H1 --year 2026

    # Arbitrary range
    python generate_report.py --from 2026-01-01 --to 2026-06-30 --label 2026-H1

Environment (.env next to this file)
------------------------------------
    ANTHROPIC_API_KEY=sk-ant-...
    HELPDESK_BASE_URL=http://localhost:5000
    HELPDESK_EMAIL=admin@gmail.com
    HELPDESK_PASSWORD=123456
    # Optional: override where the .docx is written
    # OUTPUT_DIR=./out
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import sys
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

# Headless backend MUST be set before pyplot import (cron / no display).
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

import requests
from dotenv import load_dotenv

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import anthropic


# --------------------------------------------------------------------------- #
# Config & logging
# --------------------------------------------------------------------------- #

SCRIPT_DIR = Path(__file__).resolve().parent
load_dotenv(SCRIPT_DIR / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-7s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("healthiq-report")

CLAUDE_MODEL = "claude-sonnet-4-6"
BRAND_BLUE = RGBColor(0x1F, 0x4E, 0x79)


@dataclass
class Config:
    base_url: str
    email: str
    password: str
    anthropic_api_key: str
    output_dir: Path

    @classmethod
    def from_env(cls) -> "Config":
        missing = [
            k for k in ("ANTHROPIC_API_KEY", "HELPDESK_EMAIL", "HELPDESK_PASSWORD")
            if not os.getenv(k)
        ]
        if missing:
            log.error("Missing required env vars: %s", ", ".join(missing))
            log.error("Copy .env.example to .env and fill it in.")
            sys.exit(2)

        return cls(
            base_url=os.getenv("HELPDESK_BASE_URL", "http://localhost:5000").rstrip("/"),
            email=os.environ["HELPDESK_EMAIL"],
            password=os.environ["HELPDESK_PASSWORD"],
            anthropic_api_key=os.environ["ANTHROPIC_API_KEY"],
            output_dir=Path(os.getenv("OUTPUT_DIR", str(SCRIPT_DIR / "out"))).resolve(),
        )


# --------------------------------------------------------------------------- #
# Date range helpers
# --------------------------------------------------------------------------- #

def resolve_period(args) -> tuple[str, str, str]:
    """Return (from_date, to_date, label) based on CLI args / current date."""
    if args.from_ and args.to and args.label:
        return args.from_, args.to, args.label

    if args.half and args.year:
        year = int(args.year)
        if args.half.upper() == "H1":
            return f"{year}-01-01", f"{year}-06-30", f"{year}-H1"
        if args.half.upper() == "H2":
            return f"{year}-07-01", f"{year}-12-31", f"{year}-H2"
        raise SystemExit("--half must be H1 or H2")

    today = datetime.now()
    if today.month <= 6:
        return f"{today.year}-01-01", f"{today.year}-06-30", f"{today.year}-H1"
    return f"{today.year}-07-01", f"{today.year}-12-31", f"{today.year}-H2"


# --------------------------------------------------------------------------- #
# HelpDesk API client
# --------------------------------------------------------------------------- #

class HelpDeskClient:
    def __init__(self, cfg: Config):
        self.cfg = cfg
        self.s = requests.Session()

    def login(self) -> None:
        url = f"{self.cfg.base_url}/login"
        r = self.s.post(url, json={"email": self.cfg.email, "password": self.cfg.password}, timeout=30)
        if r.status_code != 200:
            raise RuntimeError(f"Login failed ({r.status_code}): {r.text}")
        log.info("Logged in as %s", r.json().get("email"))

    def _get(self, path: str, params: dict | None = None) -> dict:
        url = f"{self.cfg.base_url}{path}"
        r = self.s.get(url, params=params or {}, timeout=60)
        if r.status_code != 200:
            raise RuntimeError(f"GET {path} failed ({r.status_code}): {r.text[:300]}")
        return r.json()

    def summary(self, from_: str, to: str) -> dict:
        return self._get("/reports/summary", {"from": from_, "to": to})

    def by_organization(self, from_: str, to: str) -> dict:
        return self._get("/reports/by-organization", {"from": from_, "to": to})

    def full_export(self, from_: str, to: str) -> dict:
        return self._get("/reports/full-export", {"from": from_, "to": to})


# --------------------------------------------------------------------------- #
# Numeric helpers
# --------------------------------------------------------------------------- #

def pct(part: int | float, whole: int | float) -> str:
    """Format part/whole as 'XX.X%' (or '-' when whole is zero)."""
    if not whole:
        return "-"
    return f"{(float(part) / float(whole) * 100.0):.1f}%"


def add_pct_column(rows: list[dict], count_key: str, total: int) -> list[dict]:
    """Return a copy of `rows` with a 'percentage' string field added."""
    out = []
    for r in rows:
        c = int(r.get(count_key) or 0)
        out.append({**r, "percentage": pct(c, total)})
    return out


# --------------------------------------------------------------------------- #
# Chart generation (matplotlib -> PNG files)
# --------------------------------------------------------------------------- #

def _savefig(path: Path) -> None:
    plt.tight_layout()
    plt.savefig(path, dpi=140, bbox_inches="tight")
    plt.close()


def chart_simple_bar(rows: list[dict], key: str, title: str, out: Path) -> None:
    """Vertical bars - good for 2-8 distinct values (severity, yes/no, etc.)."""
    labels = [(r.get(key) or "unspecified") for r in rows]
    counts = [int(r["count"]) for r in rows]
    plt.figure(figsize=(7, 3.8))
    plt.bar(labels, counts, color="#1F4E79")
    plt.title(title)
    plt.ylabel("Issues")
    plt.xticks(rotation=20, ha="right")
    plt.grid(True, axis="y", linestyle="--", alpha=0.4)
    _savefig(out)


def chart_horizontal_top(rows: list[dict], label_fn, count_key: str, title: str,
                          out: Path, top_n: int = 10) -> None:
    """Horizontal bars - good for ranked top-N (categories, orgs, indicators)."""
    top = rows[:top_n]
    labels = [label_fn(r) for r in top]
    counts = [int(r[count_key]) for r in top]
    plt.figure(figsize=(8, max(3.6, 0.4 * len(labels) + 1)))
    plt.barh(labels[::-1], counts[::-1], color="#1F4E79")
    plt.title(title)
    plt.xlabel("Issues")
    plt.grid(True, axis="x", linestyle="--", alpha=0.4)
    _savefig(out)


# --------------------------------------------------------------------------- #
# Claude analysis (structured JSON)
# --------------------------------------------------------------------------- #

ANALYSIS_PROMPT = """\
You are a data analyst preparing the analytical sections of a WHO Health IQ
helpdesk report. The report is primarily QUANTITATIVE - your job is to add
the small amount of interpretation that purely numerical sections can't
provide.

You will be given:
- Aggregate counts already computed (do NOT recompute totals or percentages)
- A list of every issue with id + description + severity + responsibility +
  category + indicator code

Return ONLY valid JSON conforming to this schema (no markdown, no prose
before or after, no code fences):

{{
  "executive_observations": [
    "string",                       // 2-3 brief, factual data observations
    "string"
  ],
  "themes": [                       // 3-5 themes total
    {{
      "theme_name": "string",       // short label, 2-5 words
      "issue_ids": [int, int, ...], // every issue id that belongs to this theme
      "explanation": "string"       // ONE short sentence
    }}
  ],
  "patterns_summary": "string",     // 1-2 sentences on concentration patterns
  "priority_areas": [               // 3-5, ordered most-to-least urgent
    {{
      "rank": 1,
      "area": "string",             // short label
      "justification": "string"     // 1 short sentence citing frequency / severity
    }}
  ],
  "recommendations": [              // 3-6, brief and practical
    {{
      "label": "string",            // 2-4 word action label
      "text": "string"              // 1 short sentence
    }}
  ]
}}

CRITICAL RULES
- Every issue must be assignable to exactly one theme. Cover all issues.
- The themes must be derived from the issue DESCRIPTIONS, not from category
  or severity values.
- Do NOT include an issue id that is not in the input.
- Do NOT mention or use the "status" field in any output.
- Be concise and factual. No marketing language.
- Prefer rounded numbers in your prose, never invent counts.

INPUT DATA:
{payload}
"""


def call_claude_for_analysis(api_key: str, payload: dict) -> dict:
    """Calls Claude with the analysis prompt, parses JSON, returns dict.
    Falls back to a minimal dict on failure so the report still builds."""
    client = anthropic.Anthropic(api_key=api_key)
    prompt = ANALYSIS_PROMPT.format(payload=json.dumps(payload, indent=2, default=str))

    log.info("Calling Claude (%s) for structured analysis...", CLAUDE_MODEL)
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        system=(
            "You output a single valid JSON object only. "
            "No prose, no commentary, no markdown code fences. "
            "Your entire reply must start with '{' and end with '}'."
        ),
        messages=[
            {"role": "user", "content": prompt},
        ],
    )
    raw = "".join(b.text for b in msg.content if getattr(b, "type", None) == "text").strip()

    # Strip ```json ... ``` fences if Claude added them despite instructions.
    if raw.startswith("```"):
        raw = raw.strip("`")
        # Drop a leading "json\n" language tag if present.
        if raw.lower().startswith("json"):
            raw = raw[4:].lstrip()

    # Slice from the first '{' to its matching '}' so any stray prose is dropped.
    raw = _trim_to_json_object(raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        log.warning("Claude returned invalid JSON (%s). Using empty fallback. Raw start: %s", e, raw[:200])
        data = {}

    return {
        "executive_observations": list(data.get("executive_observations") or []),
        "themes":                 list(data.get("themes") or []),
        "patterns_summary":       str(data.get("patterns_summary") or "").strip(),
        "priority_areas":         list(data.get("priority_areas") or []),
        "recommendations":        list(data.get("recommendations") or []),
    }


def _trim_to_json_object(s: str) -> str:
    """Return the substring from the first '{' to the matching closing '}'."""
    depth = 0
    start = s.find("{")
    if start < 0:
        return s
    for i in range(start, len(s)):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[start:i + 1]
    return s


# --------------------------------------------------------------------------- #
# DOCX low-level helpers
# --------------------------------------------------------------------------- #

def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE


def _add_para(doc: Document, text: str, *, italic: bool = False, size: int | None = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)


def _add_toc_field(doc: Document) -> None:
    """Insert a Word TOC field. User must Right-click > Update Field on first open."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    hint = doc.add_paragraph()
    hr = hint.add_run("(Right-click the table of contents above and choose 'Update Field' to populate page numbers.)")
    hr.italic = True
    hr.font.size = Pt(9)


def _add_table(doc: Document, header: list[str], rows: list[list]) -> None:
    """Render a styled table. `rows` cells are stringified."""
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = header[i]
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, "1F4E79")
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = "" if value is None else str(value)
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    for it in items:
        if it:
            doc.add_paragraph(it, style="List Bullet")


def _add_numbered_list(doc: Document, items: list[str]) -> None:
    for it in items:
        if it:
            doc.add_paragraph(it, style="List Number")


# --------------------------------------------------------------------------- #
# DOCX builder - the 8-section report
# --------------------------------------------------------------------------- #

def build_docx(
    out_path: Path,
    *,
    label: str,
    from_date: str,
    to_date: str,
    summary: dict,
    by_org: dict,
    full_export: dict,
    analysis: dict,
    charts: dict[str, Path],
) -> None:
    doc = Document()
    total = int(summary.get("total_issues", 0) or 0)

    # ---- Cover page ---------------------------------------------------------
    _cover(doc, label, from_date, to_date)
    doc.add_page_break()

    # ---- Table of contents --------------------------------------------------
    _add_heading(doc, "Table of Contents", level=1)
    _add_toc_field(doc)
    doc.add_page_break()

    # ---- 1. Executive Summary ----------------------------------------------
    _section_executive_summary(doc, total, summary, analysis)

    # ---- 2. Issue Overview --------------------------------------------------
    _section_issue_overview(doc, total, summary, charts)

    # ---- 3. Issue Distribution ---------------------------------------------
    _section_issue_distribution(doc, total, summary, by_org, charts)

    # ---- 4. Indicator-Related Issues ---------------------------------------
    _section_indicator_related(doc, total, summary, charts)

    # ---- 5. Key Patterns and Concentrations --------------------------------
    _section_patterns(doc, total, summary, by_org, analysis)

    # ---- 6. Summary of Main Issue Themes -----------------------------------
    _section_themes(doc, total, analysis)

    # ---- 7. Priority Areas --------------------------------------------------
    _section_priority(doc, analysis)

    # ---- 8. Recommendations -------------------------------------------------
    _section_recommendations(doc, analysis)

    # ---- Appendix: Raw Data -------------------------------------------------
    doc.add_page_break()
    _section_raw_appendix(doc, full_export)

    doc.save(out_path)
    log.info("DOCX saved: %s", out_path)


def _cover(doc: Document, label: str, from_date: str, to_date: str) -> None:
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_p.add_run("WHO Health IQ")
    t.bold = True
    t.font.size = Pt(32)
    t.font.color.rgb = BRAND_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Helpdesk Issue Analysis - Semi-Annual Report")
    s.font.size = Pt(16)

    doc.add_paragraph("")
    doc.add_paragraph("")

    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = period.add_run(f"Period: {from_date}  to  {to_date}  ({label})")
    pr.font.size = Pt(13)

    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gr = gen.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    gr.italic = True
    gr.font.size = Pt(11)


# ---- Section 1 -------------------------------------------------------------
def _section_executive_summary(doc: Document, total: int, summary: dict,
                                analysis: dict) -> None:
    _add_heading(doc, "1. Executive Summary", level=1)

    by_severity = summary.get("by_severity", []) or []
    by_responsibility = summary.get("by_responsibility", []) or []
    without_indicator = int(summary.get("issues_without_indicator", 0) or 0)
    with_indicator = max(0, total - without_indicator)

    doc.add_paragraph(f"Total number of issues recorded during this period: {total}")

    _add_heading(doc, "Key statistics", level=2)
    rows = []
    # Severity highlights (top 3)
    for r in by_severity[:3]:
        sev = r.get("severity") or "unspecified"
        rows.append([f"Severity = {sev}", r.get("count"), pct(r.get("count") or 0, total)])
    # Responsibility highlights (top 3)
    for r in by_responsibility[:3]:
        resp = r.get("responsibility") or "unspecified"
        rows.append([f"Responsibility = {resp}", r.get("count"), pct(r.get("count") or 0, total)])
    # Indicator-related
    rows.append(["Related to an indicator", with_indicator, pct(with_indicator, total)])
    rows.append(["NOT related to an indicator", without_indicator, pct(without_indicator, total)])
    _add_table(doc, ["Metric", "Count", "% of total"], rows)

    _add_heading(doc, "Key observations", level=2)
    obs = analysis.get("executive_observations") or []
    if obs:
        _add_bullet_list(doc, obs[:3])
    else:
        _add_para(doc, "(No observations generated.)", italic=True)


# ---- Section 2 -------------------------------------------------------------
def _section_issue_overview(doc: Document, total: int, summary: dict,
                             charts: dict[str, Path]) -> None:
    _add_heading(doc, "2. Issue Overview (Descriptive Statistics)", level=1)

    # 2a. By severity
    _add_heading(doc, "By Severity", level=2)
    sev_rows = summary.get("by_severity", []) or []
    _add_table(doc,
        ["Severity", "Count", "% of total"],
        [[r.get("severity") or "unspecified", r.get("count"), pct(r.get("count") or 0, total)]
         for r in sev_rows] or [["-", 0, "-"]],
    )
    if (p := charts.get("severity")) and p.exists():
        doc.add_picture(str(p), width=Inches(6.0))

    # 2b. By responsibility
    _add_heading(doc, "By Responsibility", level=2)
    resp_rows = summary.get("by_responsibility", []) or []
    _add_table(doc,
        ["Responsibility", "Count", "% of total"],
        [[r.get("responsibility") or "unspecified", r.get("count"), pct(r.get("count") or 0, total)]
         for r in resp_rows] or [["-", 0, "-"]],
    )
    if (p := charts.get("responsibility")) and p.exists():
        doc.add_picture(str(p), width=Inches(6.0))

    # 2c. Related to indicators (Yes/No)
    _add_heading(doc, "Related to Indicators (Yes / No)", level=2)
    without_indicator = int(summary.get("issues_without_indicator", 0) or 0)
    with_indicator = max(0, total - without_indicator)
    _add_table(doc,
        ["Related to indicator?", "Count", "% of total"],
        [
            ["Yes", with_indicator, pct(with_indicator, total)],
            ["No",  without_indicator, pct(without_indicator, total)],
        ],
    )
    if (p := charts.get("indicator_yn")) and p.exists():
        doc.add_picture(str(p), width=Inches(5.5))


# ---- Section 3 -------------------------------------------------------------
def _section_issue_distribution(doc: Document, total: int, summary: dict,
                                 by_org: dict, charts: dict[str, Path]) -> None:
    _add_heading(doc, "3. Issue Distribution", level=1)

    # 3a. By category
    _add_heading(doc, "By Category", level=2)
    cat_rows = summary.get("by_category", []) or []
    _add_table(doc,
        ["Category", "Count", "% of total"],
        [[r.get("category_name") or "uncategorized", r.get("count"), pct(r.get("count") or 0, total)]
         for r in cat_rows] or [["-", 0, "-"]],
    )
    if (p := charts.get("category")) and p.exists():
        doc.add_picture(str(p), width=Inches(6.5))

    # 3b. Top 5 categories callout
    if cat_rows:
        _add_heading(doc, "Top 5 categories", level=3)
        _add_numbered_list(doc, [
            f"{r.get('category_name') or 'uncategorized'} - {r.get('count')} issues "
            f"({pct(r.get('count') or 0, total)})"
            for r in cat_rows[:5]
        ])

    # 3c. By organization
    _add_heading(doc, "By Organization (Healthcare facility)", level=2)
    org_rows = by_org.get("organizations", []) or []
    _add_table(doc,
        ["Organization", "Count", "% of total"],
        [[o.get("organization_name") or f"Org #{o.get('organizations_id')}",
          o.get("total"), pct(o.get("total") or 0, total)]
         for o in org_rows] or [["-", 0, "-"]],
    )
    if (p := charts.get("by_org")) and p.exists():
        doc.add_picture(str(p), width=Inches(6.5))

    # 3d. Top 5 organizations callout
    if org_rows:
        _add_heading(doc, "Top 5 contributing organizations", level=3)
        top_lines = []
        for o in org_rows[:5]:
            name = o.get("organization_name") or f"Org #{o.get('organizations_id')}"
            top_lines.append(
                f"{name} - {o.get('total')} issues ({pct(o.get('total') or 0, total)})"
            )
        _add_numbered_list(doc, top_lines)


# ---- Section 4 -------------------------------------------------------------
def _section_indicator_related(doc: Document, total: int, summary: dict,
                                charts: dict[str, Path]) -> None:
    _add_heading(doc, "4. Indicator-Related Issues", level=1)

    without_indicator = int(summary.get("issues_without_indicator", 0) or 0)
    with_indicator = max(0, total - without_indicator)

    doc.add_paragraph(
        f"{with_indicator} of {total} issues are related to an indicator "
        f"({pct(with_indicator, total)}). {without_indicator} are NOT related to any indicator "
        f"({pct(without_indicator, total)})."
    )

    # 4a. Severity within indicator-related subset
    _add_heading(doc, "Severity breakdown of indicator-related issues", level=2)
    ind_sev = summary.get("indicator_severity", []) or []
    _add_table(doc,
        ["Severity", "Count", "% of indicator-related"],
        [[r.get("severity") or "unspecified", r.get("count"), pct(r.get("count") or 0, with_indicator)]
         for r in ind_sev] or [["-", 0, "-"]],
    )

    # 4b. Responsibility within indicator-related subset
    _add_heading(doc, "Responsibility breakdown of indicator-related issues", level=2)
    ind_resp = summary.get("indicator_responsibility", []) or []
    _add_table(doc,
        ["Responsibility", "Count", "% of indicator-related"],
        [[r.get("responsibility") or "unspecified", r.get("count"), pct(r.get("count") or 0, with_indicator)]
         for r in ind_resp] or [["-", 0, "-"]],
    )

    # 4c. Most common indicator-related issue types (= top indicator codes)
    _add_heading(doc, "Most common indicator-related issue types", level=2)
    ind = summary.get("by_indicator", []) or []
    if (p := charts.get("indicator")) and p.exists():
        doc.add_picture(str(p), width=Inches(6.5))
    _add_table(doc,
        ["Indicator code", "Description", "Count", "% of indicator-related"],
        [[r.get("indicator_code") or "-",
          r.get("related_to_indicators") or "-",
          r.get("count"),
          pct(r.get("count") or 0, with_indicator)]
         for r in ind] or [["-", "No indicator-linked issues", 0, "-"]],
    )


# ---- Section 5 -------------------------------------------------------------
def _section_patterns(doc: Document, total: int, summary: dict, by_org: dict,
                       analysis: dict) -> None:
    _add_heading(doc, "5. Key Patterns and Concentrations", level=1)

    cat_rows = summary.get("by_category", []) or []
    org_rows = by_org.get("organizations", []) or []
    sev_rows = summary.get("by_severity", []) or []

    # Most frequent issue types: derive from top categories
    _add_heading(doc, "Most frequent issue types", level=2)
    if cat_rows:
        _add_numbered_list(doc, [
            f"{r.get('category_name') or 'uncategorized'} - {r.get('count')} issues "
            f"({pct(r.get('count') or 0, total)})"
            for r in cat_rows[:5]
        ])
    else:
        _add_para(doc, "No category data available.", italic=True)

    # Concentration metrics
    _add_heading(doc, "Concentration metrics", level=2)
    rows = []
    if cat_rows:
        top_cat = cat_rows[0]
        rows.append([
            "Top category share",
            f"{top_cat.get('category_name') or 'uncategorized'}",
            f"{top_cat.get('count')} ({pct(top_cat.get('count') or 0, total)})",
        ])
        top3_cat = sum(int(r.get("count") or 0) for r in cat_rows[:3])
        rows.append(["Top-3 categories combined", "-", f"{top3_cat} ({pct(top3_cat, total)})"])
    if org_rows:
        top_org = org_rows[0]
        rows.append([
            "Top organization share",
            f"{top_org.get('organization_name') or 'Org #' + str(top_org.get('organizations_id'))}",
            f"{top_org.get('total')} ({pct(top_org.get('total') or 0, total)})",
        ])
        top3_org = sum(int(o.get("total") or 0) for o in org_rows[:3])
        rows.append(["Top-3 organizations combined", "-", f"{top3_org} ({pct(top3_org, total)})"])
    if sev_rows:
        top_sev = sev_rows[0]
        rows.append([
            "Most common severity",
            f"{top_sev.get('severity') or 'unspecified'}",
            f"{top_sev.get('count')} ({pct(top_sev.get('count') or 0, total)})",
        ])
    if rows:
        _add_table(doc, ["Metric", "Value", "Count (% of total)"], rows)

    # Brief AI-written summary of patterns (1-2 sentences)
    if analysis.get("patterns_summary"):
        _add_heading(doc, "Observed pattern", level=2)
        doc.add_paragraph(analysis["patterns_summary"])


# ---- Section 6 -------------------------------------------------------------
def _section_themes(doc: Document, total: int, analysis: dict) -> None:
    _add_heading(doc, "6. Summary of Main Issue Themes", level=1)

    themes = analysis.get("themes") or []
    if not themes:
        _add_para(doc, "(No themes generated.)", italic=True)
        return

    # Resolve issue_ids -> count + percentage. Claude provides ids; we own the math.
    rows = []
    for t in themes:
        ids = t.get("issue_ids") or []
        cnt = len(ids)
        rows.append([
            t.get("theme_name") or "(unnamed theme)",
            cnt,
            pct(cnt, total),
            t.get("explanation") or "",
        ])
    _add_table(doc,
        ["Theme", "Issues", "% of total", "Explanation"],
        rows,
    )


# ---- Section 7 -------------------------------------------------------------
def _section_priority(doc: Document, analysis: dict) -> None:
    _add_heading(doc, "7. Priority Areas", level=1)

    pri = analysis.get("priority_areas") or []
    if not pri:
        _add_para(doc, "(No priority areas generated.)", italic=True)
        return

    # Render as a numbered ranked list
    pri_sorted = sorted(pri, key=lambda x: int(x.get("rank") or 999))
    _add_numbered_list(doc, [
        f"{p.get('area') or '(unnamed)'} - {p.get('justification') or ''}"
        for p in pri_sorted
    ])


# ---- Section 8 -------------------------------------------------------------
def _section_recommendations(doc: Document, analysis: dict) -> None:
    _add_heading(doc, "8. Recommendations", level=1)

    recs = analysis.get("recommendations") or []
    if not recs:
        _add_para(doc, "(No recommendations generated.)", italic=True)
        return

    for r in recs:
        p = doc.add_paragraph(style="List Bullet")
        label = (r.get("label") or "").strip()
        text = (r.get("text") or "").strip()
        if label:
            run = p.add_run(f"{label}: ")
            run.bold = True
        p.add_run(text)


# ---- Appendix --------------------------------------------------------------
def _section_raw_appendix(doc: Document, full_export: dict) -> None:
    _add_heading(doc, "Appendix: Raw Data", level=1)
    _add_para(doc, f"{full_export.get('count', 0)} issues included in the full export.", italic=True)
    issues = full_export.get("issues", []) or []
    header = ["ID", "Start Date", "Org", "Category", "Severity", "Indicator", "Description"]
    rows = []
    for it in issues:
        start_raw = it.get("startDate") or ""
        rows.append([
            it.get("id"),
            str(start_raw)[:10],
            it.get("organization_name") or "-",
            it.get("category_name") or "-",
            it.get("severity") or "-",
            it.get("indicator_code") or "-",
            (it.get("description") or "")[:180],
        ])
    _add_table(doc, header, rows)


# --------------------------------------------------------------------------- #
# Post-generation hooks (stubbed per current project config)
# --------------------------------------------------------------------------- #

def upload_to_google_drive(docx_path: Path) -> None:
    log.info("[skip] Google Drive upload is disabled in this build.")


def send_email_notification(docx_path: Path, label: str) -> None:
    log.info("[skip] Email notification is disabled in this build.")


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #

def parse_args(argv: list[str]) -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate the WHO Health IQ helpdesk report.")
    p.add_argument("--from", dest="from_", help="Start date YYYY-MM-DD")
    p.add_argument("--to", help="End date YYYY-MM-DD")
    p.add_argument("--label", help="Period label, e.g. '2026-H1'")
    p.add_argument("--half", choices=["H1", "H2", "h1", "h2"], help="Half-year")
    p.add_argument("--year", type=int, help="Year for --half")
    return p.parse_args(argv)


def _build_claude_payload(summary: dict, by_org: dict, full_export: dict,
                           label: str, from_date: str, to_date: str) -> dict:
    """Compact payload for Claude. Pre-aggregated counts + raw issue
    descriptions (the only thing Claude needs the raw data for, for theming)."""
    issues = full_export.get("issues", []) or []
    issue_brief = [
        {
            "id": it.get("id"),
            "description": (it.get("description") or "")[:400],
            "severity": it.get("severity"),
            "responsibility": it.get("responsibility"),
            "category": it.get("category_name"),
            "indicator_code": it.get("indicator_code"),
        }
        for it in issues
    ]
    return {
        "period": {"from": from_date, "to": to_date, "label": label},
        "totals": {
            "total_issues": summary.get("total_issues"),
            "issues_without_indicator": summary.get("issues_without_indicator"),
        },
        "by_severity":              summary.get("by_severity"),
        "by_responsibility":        summary.get("by_responsibility"),
        "by_category":              summary.get("by_category"),
        "by_role":                  summary.get("by_role"),
        "by_indicator":             summary.get("by_indicator"),
        "indicator_severity":       summary.get("indicator_severity"),
        "indicator_responsibility": summary.get("indicator_responsibility"),
        "by_organization":          by_org.get("organizations"),
        "issues":                   issue_brief,
    }


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    cfg = Config.from_env()
    cfg.output_dir.mkdir(parents=True, exist_ok=True)

    from_date, to_date, label = resolve_period(args)
    log.info("Report period: %s  ->  %s   [label=%s]", from_date, to_date, label)

    # 1. Pull data
    client = HelpDeskClient(cfg)
    client.login()
    summary     = client.summary(from_date, to_date)
    by_org      = client.by_organization(from_date, to_date)
    full_export = client.full_export(from_date, to_date)
    total = int(summary.get("total_issues", 0) or 0)
    log.info("Pulled %d issues", full_export.get("count", 0))

    if total == 0:
        log.warning("No issues found for this period - report will be mostly empty.")

    # 2. Generate charts (only ones the new structure needs)
    with tempfile.TemporaryDirectory(prefix="healthiq-charts-") as tmp:
        tmp_dir = Path(tmp)
        charts: dict[str, Path] = {}

        if summary.get("by_severity"):
            charts["severity"] = tmp_dir / "severity.png"
            chart_simple_bar(summary["by_severity"], "severity",
                             "Issues by severity", charts["severity"])

        if summary.get("by_responsibility"):
            charts["responsibility"] = tmp_dir / "responsibility.png"
            chart_simple_bar(summary["by_responsibility"], "responsibility",
                             "Issues by responsibility", charts["responsibility"])

        # Yes / No indicator chart
        without = int(summary.get("issues_without_indicator", 0) or 0)
        if total > 0:
            charts["indicator_yn"] = tmp_dir / "indicator_yn.png"
            chart_simple_bar(
                [{"label": "Yes", "count": max(0, total - without)},
                 {"label": "No",  "count": without}],
                "label", "Issues related to an indicator", charts["indicator_yn"],
            )

        if summary.get("by_category"):
            charts["category"] = tmp_dir / "category.png"
            chart_horizontal_top(
                summary["by_category"],
                label_fn=lambda r: (r.get("category_name") or "uncategorized")[:40],
                count_key="count",
                title="Top categories by issue count",
                out=charts["category"], top_n=10,
            )

        if by_org.get("organizations"):
            charts["by_org"] = tmp_dir / "by_org.png"
            chart_horizontal_top(
                by_org["organizations"],
                label_fn=lambda o: (o.get("organization_name")
                                    or f"Org #{o.get('organizations_id')}")[:40],
                count_key="total",
                title="Top organizations by issue count",
                out=charts["by_org"], top_n=10,
            )

        if summary.get("by_indicator"):
            charts["indicator"] = tmp_dir / "indicator.png"
            chart_horizontal_top(
                summary["by_indicator"],
                label_fn=lambda r: (
                    f"{r.get('indicator_code') or '?'}"
                    + (f" - {(r.get('related_to_indicators') or '')[:30]}"
                       if r.get('related_to_indicators') else "")
                ),
                count_key="count",
                title="Top indicators by issue count",
                out=charts["indicator"], top_n=12,
            )

        # 3. Ask Claude for the structured analysis
        payload = _build_claude_payload(summary, by_org, full_export,
                                         label, from_date, to_date)
        analysis = call_claude_for_analysis(cfg.anthropic_api_key, payload)
        log.info("Claude returned %d themes, %d priority areas, %d recommendations",
                 len(analysis["themes"]), len(analysis["priority_areas"]),
                 len(analysis["recommendations"]))

        # 4. Build the DOCX
        docx_path = cfg.output_dir / f"HealthIQ_Report_{label}.docx"
        build_docx(
            docx_path,
            label=label,
            from_date=from_date, to_date=to_date,
            summary=summary, by_org=by_org, full_export=full_export,
            analysis=analysis, charts=charts,
        )

    # 5. Post-generation hooks (currently stubbed)
    upload_to_google_drive(docx_path)
    send_email_notification(docx_path, label)

    log.info("Done. Report: %s", docx_path)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main(sys.argv[1:]))
    except RuntimeError as e:
        log.error("%s", e)
        raise SystemExit(1)
